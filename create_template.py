
import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_approval_template(path):
    doc = Document()
    
    # Titolo
    title = doc.add_heading("Verbale di Approvazione del Bilancio", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Ente e Intestazione
    p = doc.add_paragraph()
    p.add_run("Associazione: ").bold = True
    p.add_run("{{ ente }}")
    
    p = doc.add_paragraph()
    p.add_run("Data: ").bold = True
    p.add_run("{{ data_oggi }}")
    
    doc.add_paragraph("\nIn data odierna si è riunito l'organo competente per l'approvazione del rendiconto per cassa.")
    
    # Dati Bilancio
    doc.add_heading("Sintesi del Rendiconto {{ anno_consuntivo }}", level=1)
    
    table_data = [
        ["Voce", "Importo (€)"],
        ["Totale Entrate", "{{ totale_entrate }}"],
        ["Totale Uscite", "{{ totale_uscite }}"],
        ["Avanzo/Disavanzo d'Esercizio", "{{ avanzo }}"],
        ["Avanzo dopo le imposte", "{{ avanzo_finale }}"],
    ]
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = table_data[0][0]
    hdr_cells[1].text = table_data[0][1]
    
    for voce, valore in table_data[1:]:
        row_cells = table.add_row().cells
        row_cells[0].text = voce
        row_cells[1].text = valore
        
    doc.add_paragraph("\n{{ frase_saldo }}\n")
    
    # Saldi finali
    doc.add_heading("Situazione Patrimoniale al 31/12/{{ anno_consuntivo }}", level=2)
    p = doc.add_paragraph()
    p.add_run("Saldo Cassa: ").bold = True
    p.add_run("€ {{ saldo_cassa }}")
    p = doc.add_paragraph()
    p.add_run("Saldo Conto Corrente: ").bold = True
    p.add_run("€ {{ saldo_cc }}")
    p = doc.add_paragraph()
    p.add_run("TOTALE GENERALE: ").bold = True
    p.add_run("€ {{ totale_generale }}")
    
    doc.add_paragraph("\nIl presente verbale viene letto, confermato e sottoscritto.")
    
    # Firme
    doc.add_paragraph("\n\n")
    firme_para = doc.add_paragraph()
    firme_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    firme_para.add_run("Il Presidente\t\tIl Segretario\n\n\n_______________________\t\t_______________________")
    
    doc.save(path)
    print(f"Template creato con successo in: {path}")

if __name__ == "__main__":
    template_path = "template_verbale_approvazione.docx"
    create_approval_template(template_path)
